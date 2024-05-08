from docx import Document
from docx.shared import Pt
import re

# Load your docx file
input_filename = "input.docx"  # Replace with your actual file path
output_filename = "output4.docx"
document = Document(input_filename)

# Regular expressions to find bold text and lines starting with ###
bold_pattern = re.compile(r"\*\*(.*?)\*\*")
heading_pattern = re.compile(r"^###\s*(.*)")

# Iterate through paragraphs and apply styles
for para in document.paragraphs:
    text = para.text

    # Check for lines starting with '###'
    heading_match = heading_pattern.match(text)
    if heading_match:
        # Extract heading text and clear paragraph
        heading_text = heading_match.group(1)
        para.clear()
        run = para.add_run(heading_text)
        run.font.size = Pt(14)  # Set the heading size to 14 points
        run.bold = True  # Make the heading text bold
        continue

    # Check for bold text enclosed in ** **
    bold_matches = bold_pattern.findall(text)
    if bold_matches:
        # Split the paragraph text based on the bold pattern
        parts = bold_pattern.split(text)
        para.clear()  # Clear the existing paragraph

        # Iterate over split parts, applying bold to matches
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Odd index indicates it's between double asterisks
                run = para.add_run(part)
                run.bold = True
            else:
                para.add_run(part)

    # else:
    #     # No specific matches, retain the original paragraph text
    #     para.text = text

# Save the updated document
document.save(output_filename)
